import os
from PyPDF2 import PdfReader, PdfWriter, PdfMerger
from PIL import Image
from deep_translator import GoogleTranslator
from pdf2docx import Converter
from PIL import Image
from pytesseract import pytesseract
import docx
import re


def handle_uploaded_file(f):
    """save some pdf doc to translate."""
    with open("some/file/name.pdf", "wb+") as destination:
        for chunk in f.chunks():
            destination.write(chunk)


def extract_images_from_pdf(pdf_path):
    """extract image from pdf file."""
    with open(pdf_path, "rb") as f:
        reader = PdfReader(f)
        for page_num in enumerate(len(reader.pages)):
            selected_page = reader.pages[page_num]
            for img_file_obj in selected_page.images:
                with open(img_file_obj.name, "wb") as out:
                    out.write(img_file_obj.data)


def split_pdf(pdf_path):
    """."""
    # Function to Split PDF into Multiple PDF Pages
    with open(pdf_path, "rb") as f:
        reader = PdfReader(f)
        # get all pages
        for page_num in enumerate(len(reader.pages)):  # loop through pages
            selected_page = reader.pages[page_num]
            # Writer to write
            writer = PdfWriter()
            writer.add_page(selected_page)  # add/embedding of the page
            filename = os.path.splitext(pdf_path)[0]
            output_filename = f"{filename}_page_{page_num + 1}.pdf"
            # save and compile to pdf
            with open(output_filename, "wb") as out:
                writer.write(out)

            print("created a pdf:{}".format(output_filename))


def get_last_pdf_page(pdf_path):
    """."""
    # GET LAST Pages IN PDF FILE
    with open(pdf_path, "rb") as f:
        reader = PdfReader(f)
        writer = PdfWriter()
        selected_page = reader.pages[len(reader.pages) - 1]
        writer.add_page(selected_page)
        filename = os.path.splitext(pdf_path)[0]
        output_filename = f"{filename}_last_page.pdf"
        with open(output_filename, "wb") as out:
            writer.write(out)
        print("created last page")


# Merging PDFs
# get a list of pdfs


def fetch_all_pdf_files(parent_folder: str):
    # Pdf Fetch
    """."""
    target_files = []
    for path, files in os.walk(parent_folder):
        for name in files:
            if name.endswith(".pdf"):
                target_files.append(os.path.join(path, name))
    return target_files


def merge_pdf(list_of_pdfs, output_filename="final_merged_file.pdf"):
    """."""
    # PdfMerger
    merger = PdfMerger()
    with open(output_filename, "wb") as f:
        for file in list_of_pdfs:
            merger.append(file)
        merger.write(f)


def convert_img2pdf(image_file):
    """."""
    # Set-up and wrap translation client
    my_image = Image.open(image_file)
    img = my_image.convert("RGB")
    filename = f"{os.path.splitext(image_file)[0]}.pdf"
    img.save(filename)


def extract_text_from_pdf(pdf_path):
    """."""
    with open(pdf_path, "rb") as f:
        reader = PdfReader(f)
        results = []
        for i in enumerate(len(reader.pages)):  # prev read.getNumPages()
            selected_page = reader.pages[i]
            text = selected_page.extract_text()
            results.append(text)
    return " ".join(results)  # convert list to a single doc


def translate_extracted(text, lng="fr"):
    """."""
    # Set-up and wrap translation client
    translate = GoogleTranslator(source="auto", target=lng).translate
    #     #Split input text into a list of sentences
    #     sentences = sent_tokenize(Extracted)
    #     # Initialize containers
    #     translated_text = ''
    #     source_text_chunk = ''
    #     # collect chuncks of sentences, translate individually
    #     for sentence in sentences:
    #       # if chunck + current sentence < limit, add the sentence
    #       if ((len(sentence.encode('utf-8')) +  len(source_text_chunk.encode('utf-8')) < 5000)):
    #         source_text_chunk += ' ' + sentence
    #       # else translate chunck and start new one with current sentence
    #       else:
    #         translated_text += ' ' + translate(source_text_chunk)
    #       # if current sentence smaller than 5000 chars, start new chunck
    #       if (len(sentence.encode('utf-8')) < 5000):
    #             source_text_chunk = sentence
    #       # else, replace sentence with notification message
    #       else:
    #             message = "<<Omitted Word longer than 5000bytes>>"
    #             translated_text += ' ' + translate(message)

    #             # Re-set text container to empty
    #             source_text_chunk = ''
    #      # Translate the final chunk of input text, if there is any valid text left to translate
    #       if translate(source_text_chunk) != None:
    #             translated_text += ' ' + translate(source_text_chunk)

    #       return translated_text
    return translate(text)


def read_file_line_by_line(path):
    """Function to read file line by line."""
    with open(path, "rb") as f:
        pdf_reader = PdfReader(f)
        for page in pdf_reader.pages:
            for line in page.extract_text().splitlines():
                yield line


def read_text_line_by_line(texte):
    """Function to read file line by line."""

    for line in texte.splitlines():
        yield line


def translate_pdf_file(path):
    """Function to translate pdf."""
    result = []
    return pdf_to_docx(path)
    # with open(path, "rb") as f:
    #     reader = PdfReader(f)
    #     # get all pages
    #     print("start work")
    #     for page in reader.pages:  # loop through pages
    #         text = page.extract_text()
    #         for line in read_text_line_by_line(text):
    #             res = translate_extracted(line)
    #             print(res)
    #             res = res + "\n"
    #             result.append(res)

    #     filename = os.path.splitext(path)[0]
    #     output_filename = f"{filename}_translate.doc"
    #     with open(output_filename, "wb") as destination:
    #         for chunk in result:
    #             destination.write(chunk.encode())

    #     return output_filename


def pdf_to_docx(path):
    """Function to convert pdf to docx."""
    c = Converter(path)
    filename = os.path.splitext(path)[0]
    output_filename = f"{filename}_.docx"
    c.convert(output_filename)
    c.close()
    # text = docx2txt.process(output_filename)
    # print(text)
    doc = docx.Document(output_filename)
    all_paras = doc.paragraphs
    for para in all_paras:
        if special_match(para.text):
            para.text = translate_extracted(para.text)
            print(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if special_match(paragraph.text):
                        paragraph.text = translate_extracted(paragraph.text)
    file = os.path.splitext(path)[0]
    output_save = f"{file}_translate.doc"
    doc.save(output_save)
    return output_save
    # time.sleep(20000)


def special_match(strg, search=re.compile(r"[^a-z0-9.]").search):
    """Function to check if paragraph contain some text."""
    return bool(search(strg))


# def special_match(image_path):
#     """Extract texte from an image."""
#     return bool(search(strg))
